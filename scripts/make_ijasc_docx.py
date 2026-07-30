"""Assemble docs/ijasc_manuscript.md into an IJASC-formatted .docx.

Format rules implemented from docs/ijasc_format_notes.md:
  page 21x28 cm, margins 2 cm sides / 3 cm top / 2 cm bottom (text block 17x23 cm)
  title 15 pt TNR bold italic centred, spacing 1.1
  author block 11.5 pt bold centred; affiliation 10.5 pt centred
  "Abstract" 11 pt bold centred; abstract body 9 pt, spacing 1.15
  keywords 9 pt, label bold, indented
  body 9.5 pt TNR, spacing 1.15, justified, first-line indent ~1 character
  table captions ABOVE (9 pt centred), figure captions BELOW (9 pt centred)
  "REFERENCES" 11 pt bold centred; entries 9 pt, hanging indent 0.8 cm
  single column throughout (both IJASC sample papers are single column)

Output: docs/ijasc_manuscript.docx
"""
import re, os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
MD   = os.path.join(ROOT, "docs", "ijasc_manuscript.md")
OUT  = os.path.join(ROOT, "docs", "ijasc_manuscript.docx")
FIGD = os.path.join(ROOT, "data", "paper_figures")

FIGMAP = {1: "fig1_mechanism.png",          2: "fig1_pipeline_predicates.png",
          3: "fig3_stall_and_engagement.png", 4: "fig2_outcome_divergence.png",
          5: "fig4_engagement_conditions.png", 6: "fig6_interventions.png"}

TEXT_W = Cm(17.0)

# --------------------------------------------------------------- inline runs
TOKEN = re.compile(r"(\*\*.+?\*\*|\*[^*]+\*)")
SUB   = re.compile(r"_([A-Za-z][A-Za-z0-9]*|box|0)")

FORBID = " ([{|-–—\u2016"

def add_runs(par, text, base_size, force_italic=False):
    """Markdown *italic* / **bold**; attached _x becomes a true subscript."""
    text = (text.replace(r"\|", "|")
                .replace(r"\*", "\x01").replace(r"\_", "\x02"))
    last = ""                                  # last character emitted, across pieces
    for piece in TOKEN.split(text):
        if not piece:
            continue
        bold = italic = False
        if piece.startswith("**") and piece.endswith("**"):
            bold, piece = True, piece[2:-2]
        elif piece.startswith("*") and piece.endswith("*"):
            italic, piece = True, piece[1:-1]
        piece = piece.replace("\x01", "*").replace("\x02", "_")
        pos = 0
        for m in SUB.finditer(piece):
            prev = piece[m.start()-1] if m.start() > 0 else last
            if not prev or prev in FORBID:
                continue                       # "_" not attached to a symbol
            if m.start() > pos:
                emit_with_unicode_subs(par, piece[pos:m.start()], base_size, bold,
                                       italic or force_italic)
            r = par.add_run(m.group(1))
            r.font.size, r.bold = base_size, bold
            r.italic = italic or force_italic
            r.font.subscript = True
            pos = m.end()
        if pos < len(piece):
            emit_with_unicode_subs(par, piece[pos:], base_size, bold,
                                   italic or force_italic)
        if piece:
            last = piece[-1]


USUB = {"\u2080":"0","\u2081":"1","\u2082":"2","\u2083":"3","\u2084":"4",
        "\u2085":"5","\u2086":"6","\u2087":"7","\u2088":"8","\u2089":"9",
        "\u208a":"+","\u208b":"-","\u2096":"k"}

def emit_with_unicode_subs(par, text, base_size, bold, italic):
    """Split text on unicode-subscript characters and emit them as true subscripts."""
    buf = ""
    def flush_plain():
        nonlocal buf
        if buf:
            r = par.add_run(buf)
            r.font.size, r.bold, r.italic = base_size, bold, italic
            buf = ""
    subbuf = ""
    def flush_sub():
        nonlocal subbuf
        if subbuf:
            r = par.add_run(subbuf)
            r.font.size, r.bold, r.italic = base_size, bold, italic
            r.font.subscript = True
            subbuf = ""
    for ch in text:
        if ch in USUB:
            flush_plain(); subbuf += USUB[ch]
        else:
            flush_sub(); buf += ch
    flush_plain(); flush_sub()

def keep_next(par):
    par.paragraph_format.keep_with_next = True
    return par

def row_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:cantSplit"); trPr.append(el)

def P(doc, size=Pt(9.5), align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=Cm(0.35),
      before=Pt(0), after=Pt(0), line=1.15):
    par = doc.add_paragraph()
    pf = par.paragraph_format
    pf.alignment = align
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line
    pf.space_before, pf.space_after = before, after
    if indent:
        pf.first_line_indent = indent
    return par

# --------------------------------------------------------------- md parsing
raw = open(MD, encoding="utf-8").read()
raw = raw.split("---", 1)[1]                     # drop the assembly header
lines = raw.split("\n")

doc = Document()
sec = doc.sections[0]
sec.page_width,  sec.page_height  = Cm(21.0), Cm(28.0)
sec.left_margin, sec.right_margin = Cm(2.0),  Cm(2.0)
sec.top_margin,  sec.bottom_margin = Cm(3.0), Cm(2.0)
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
style.font.size = Pt(9.5)

i, n = 0, len(lines)
mode = None            # None | 'abstract' | 'refs'
para_buf = []

def flush_para():
    global para_buf
    if not para_buf:
        return
    text = " ".join(x.strip() for x in para_buf)
    para_buf = []
    if not text:
        return
    if mode == "abstract":
        par = P(doc, size=Pt(9), indent=None)
        add_runs(par, text, Pt(9))
    elif text.startswith("**Keywords:**"):
        par = P(doc, size=Pt(9), align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=None,
                before=Pt(6), line=1.1)
        par.paragraph_format.left_indent = Cm(0.7)
        r = par.add_run("Keywords: "); r.bold = True; r.font.size = Pt(9)
        add_runs(par, text[len("**Keywords:**"):].strip(), Pt(9))
    else:
        par = P(doc)
        add_runs(par, text, Pt(9.5))

def caption_par(text, size=Pt(9)):
    par = P(doc, size=size, align=WD_ALIGN_PARAGRAPH.CENTER, indent=None,
            before=Pt(4), after=Pt(6))
    par.paragraph_format.keep_together = True          # never split a caption
    add_runs(par, text, size)
    return par

def add_table(cap_text, rows):
    keep_next(caption_par(cap_text))                    # caption ABOVE, tied to table
    cells = [ [c.strip() for c in re.split(r"(?<!\\)\|", r)[1:-1] ] for r in rows]
    ncol = len(cells[0])
    t = doc.add_table(rows=len(cells), cols=ncol)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    # proportional column widths over the 17 cm block, floor 1.5 cm
    maxlen = [max(len(cells[ri][ci]) for ri in range(len(cells))) ** 0.8
              for ci in range(ncol)]
    total = sum(maxlen)
    widths = [max(1.5, 17.0 * m / total) for m in maxlen]
    scale = 17.0 / sum(widths)
    widths = [w * scale for w in widths]
    for ri, row in enumerate(cells):
        row_cant_split(t.rows[ri])
        for ci, cell in enumerate(row):
            c = t.cell(ri, ci)
            c.width = Cm(widths[ci])
            c.paragraphs[0].text = ""
            par = c.paragraphs[0]
            par.paragraph_format.line_spacing = 1.0
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER if ri == 0 else WD_ALIGN_PARAGRAPH.LEFT
            if ri < len(cells) - 1:
                par.paragraph_format.keep_with_next = True
            add_runs(par, cell, Pt(9))
            if ri == 0:
                for r in par.runs: r.bold = True
    spacer = P(doc, indent=None, after=Pt(6)); spacer.add_run("")

FIGW = {1: Cm(16.0), 3: Cm(15.0), 6: Cm(15.0)}                      # tall figures slightly narrower

def add_figure(num, cap_text):
    path = os.path.join(FIGD, FIGMAP[num])
    par = P(doc, align=WD_ALIGN_PARAGRAPH.CENTER, indent=None, before=Pt(6))
    par.add_run().add_picture(path, width=FIGW.get(num, TEXT_W))
    keep_next(par)                                      # image stays with caption
    caption_par(cap_text)                               # caption BELOW

while i < n:
    line = lines[i]
    st = line.strip()

    if st.startswith("## Title"):
        flush_para(); i += 1
        while not lines[i].strip(): i += 1
        title = lines[i].strip().strip("*")
        par = P(doc, align=WD_ALIGN_PARAGRAPH.CENTER, indent=None, line=1.1,
                after=Pt(11))
        r = par.add_run(title); r.font.size = Pt(15); r.bold = True; r.italic = True
        i += 1; continue

    if st.startswith("## Authors"):
        flush_para(); i += 1
        par = P(doc, align=WD_ALIGN_PARAGRAPH.CENTER, indent=None, line=1.8,
                before=Pt(11))
        r = par.add_run("Author Name"); r.font.size = Pt(11.5); r.bold = True
        par = P(doc, align=WD_ALIGN_PARAGRAPH.CENTER, indent=None, after=Pt(10))
        r = par.add_run("Affiliation, Department, University, Country  (fill per template)")
        r.font.size = Pt(10.5)
        while i < n and not lines[i].strip().startswith("## "): i += 1
        continue

    if st.startswith("## Abstract"):
        flush_para()
        par = P(doc, align=WD_ALIGN_PARAGRAPH.CENTER, indent=None, after=Pt(7))
        r = par.add_run("Abstract"); r.font.size = Pt(11); r.bold = True
        mode = "abstract"; i += 1; continue

    if st.startswith("**Keywords:**"):
        flush_para(); mode = None
        para_buf.append(st)
        i += 1
        while i < n and lines[i].strip() and not lines[i].startswith("#"):
            para_buf.append(lines[i]); i += 1
        flush_para(); continue

    if st.startswith("## References"):
        flush_para(); mode = "refs"
        par = P(doc, align=WD_ALIGN_PARAGRAPH.CENTER, indent=None,
                before=Pt(12), after=Pt(6))
        keep_next(par)
        r = par.add_run("REFERENCES"); r.font.size = Pt(11); r.bold = True
        i += 1; continue

    if st.startswith("## Acknowledgement"):
        flush_para(); mode = None
        par = P(doc, align=WD_ALIGN_PARAGRAPH.LEFT, indent=None,
                before=Pt(12), after=Pt(4))
        r = par.add_run("Acknowledgement"); r.font.size = Pt(11); r.bold = True
        par = P(doc, size=Pt(9.5))
        par.add_run("This work was supported by (institution to acknowledge) in (year).").font.size = Pt(9.5)
        i += 1
        while i < n and not lines[i].strip().startswith("## "): i += 1
        continue

    if re.match(r"^## \d", st):                       # numbered section
        flush_para(); mode = None
        par = P(doc, align=WD_ALIGN_PARAGRAPH.LEFT, indent=None,
                before=Pt(12), after=Pt(6), line=1.1)
        keep_next(par)
        r = par.add_run(st[3:].strip()); r.font.size = Pt(11); r.bold = True
        i += 1; continue

    if re.match(r"^### \d", st):                      # subsection
        flush_para()
        par = P(doc, align=WD_ALIGN_PARAGRAPH.LEFT, indent=None,
                before=Pt(10), after=Pt(4), line=1.1)
        keep_next(par)
        r = par.add_run(st[4:].strip()); r.font.size = Pt(10); r.bold = True
        i += 1; continue

    if st == "---" or st.startswith("## "):
        flush_para(); i += 1; continue

    m = re.match(r"^\*\*\[Figure (\d) here\]\*\*", st)
    if m:
        flush_para()
        num = int(m.group(1)); i += 1
        while i < n and not lines[i].strip(): i += 1
        cap = []
        while i < n and lines[i].strip():
            cap.append(lines[i].strip()); i += 1
        cap_text = " ".join(cap).strip("*").strip()
        add_figure(num, cap_text)
        continue

    if st.startswith("*Table ") or st.startswith("*Table"):
        flush_para()
        cap = [st]
        i += 1
        while i < n and lines[i].strip() and not lines[i].strip().startswith("|"):
            cap.append(lines[i].strip()); i += 1
        cap_text = " ".join(cap).strip("*").strip()
        while i < n and not lines[i].strip(): i += 1
        rows = []
        while i < n and lines[i].strip().startswith("|"):
            if not re.match(r"^\|[\s:-]+\|", lines[i].strip().replace("-", "-")) \
               or "---" not in lines[i]:
                rows.append(lines[i].strip())
            i += 1
        rows = [r for r in rows if "---" not in r]
        add_table(cap_text, rows)
        continue

    if re.search(r"\(\d\)\s*$", st) and not st.startswith(("*Table", "|")) \
       and len(st) < 120 and ("=" in st or "∧" in st):
        flush_para()                                   # display equation
        par = P(doc, align=WD_ALIGN_PARAGRAPH.CENTER, indent=None,
                before=Pt(6), after=Pt(6))
        add_runs(par, st, Pt(9.5))
        i += 1; continue

    if st.startswith("* "):                            # bullet
        flush_para()
        par = P(doc, indent=None, line=1.15)
        par.paragraph_format.left_indent = Cm(0.6)
        par.paragraph_format.first_line_indent = Cm(-0.35)
        add_runs(par, "•  " + st[2:], Pt(9.5))
        i += 1
        while i < n and lines[i].startswith("  ") and lines[i].strip():
            add_runs(par, " " + lines[i].strip(), Pt(9.5)); i += 1
        continue

    if re.match(r"^\d\. ", st):                        # numbered list item
        flush_para()
        par = P(doc, indent=None)
        par.paragraph_format.left_indent = Cm(0.6)
        par.paragraph_format.first_line_indent = Cm(-0.35)
        add_runs(par, st, Pt(9.5))
        i += 1
        while i < n and lines[i].startswith("   ") and lines[i].strip():
            add_runs(par, " " + lines[i].strip(), Pt(9.5)); i += 1
        continue

    if mode == "refs" and re.match(r"^\[\d+\]", st):
        entry = [st]; i += 1
        while i < n and lines[i].strip() and not re.match(r"^\[\d+\]", lines[i].strip()):
            entry.append(lines[i].strip()); i += 1
        par = P(doc, size=Pt(9), indent=None)
        par.paragraph_format.left_indent = Cm(0.8)
        par.paragraph_format.first_line_indent = Cm(-0.8)
        add_runs(par, " ".join(entry), Pt(9))
        continue

    if not st:
        flush_para(); i += 1; continue

    para_buf.append(line); i += 1

flush_para()
doc.save(OUT)
print("wrote", OUT)
